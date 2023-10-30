#!/usr/bin/env python3
"""
Extraction functions for each file type to be used with Document class

Notes on getting text:
* assume docs are max 100 pages - only get N bytes of string to conserve memory
* `text_lst` itesms should be in batches, such as doc pages or paragraphs (don't mess up sentenceizer)
* `docs = spacy.pipe(text_lst, n_process=-1)`   #as many processes as CPUs
* resulting docs is a generator, but return as list `list(docs)`, currently
* then Document performs text processing

TODO:minor tasks
* convert all file types to pdf, then extract using pdf.miner
* limit retrieval to N excerpts
* use optimal storage: `from collections import namedtuple`   #ref: https://stackoverflow.com/questions/1336791/dictionary-vs-object-which-is-more-efficient-and-why

#TODO:review PyMuPDf
#PyMuPDF is AGPL
#PyMuPDF is faster than pdfminer, [ref](https://medium.com/social-impact-analytics/comparing-4-methods-for-pdf-text-extraction-in-python-fd34531034f)
#PyMuPDF extracts ToC if the PDF consists of Bookmarks. Otherwise it only results in an empty list.
#PyMuPDF appears to extract ToC better, but look into open-license methods: https://stackoverflow.com/questions/54303318/read-all-bookmarks-from-a-pdf-document-and-create-a-dictionary-with-pagenumber-a
#pikepdf may be an alternative for getting ToC => it gets bookmarks, but sometimes not page_location

#TODO:add ocr
#OCRmyPDF uses pikepdf, but is MPL-2.0 license (pikepdf also), which requires commercial sales to be open-sourced if files are distributed
#in server SAAS solution, no files are distributed - so probably good [ref](https://opensource.stackexchange.com/questions/9407/do-i-have-to-disclose-changes-made-to-mpl-2-0-licensed-code-running-on-server#:~:text=For%20server-side%20software%20offered%20in%20a%20SaaS%20setting%2C,the%20source%20code%20available%20alongside%20any%20binary%20distribution.)
#madmaze/pytesseract is Apache-2.0 license
"""

__author__ = "Jason Beach"
__version__ = "0.1.0"
__license__ = "MIT"

from .url import UniformResourceLocator
from .record import DocumentRecord
from .config import EnteroConfig
'''
from .config import (
    applySpacy,
    applyPyMuPDF
)
from .utils import (
    MAX_PAGE_EXTRACT,
    timeout
)'''

#html
import bs4
from bs4.element import Comment
import xhtml2pdf
from xhtml2pdf import pisa 

#pdf
import pypdf
import pdftitle                                                                 #uses pdfminer
from pdfminer.high_level import extract_text as pdf_extract_text                #uses pdfminer.six, problem TODO???
from pdfminer.high_level import extract_pages as pdf_extract_pages
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import resolve1

import fitz
from pikepdf import Pdf



#docx
import docx


import pandas as pd     #TODO:replace logic to remove module

import datetime
import io
import time
from pathlib import Path, PosixPath




docrec = DocumentRecord()
record = docrec._asdict()





def get_title(txt):
    pass

def clean_text(txt):
    if type(txt) is list:
        combined_txt = ('.').join(txt)
        return combined_txt.replace('.','.  ').replace('\n',' ')
    elif type(txt) is str:
        txts = txt.split('.\n')
        #TODO:if the len(item)<50, then append to the earlier item
        txts = [txt.replace('-\n','').replace('\n',' ') for txt in txts]
        return txts
    else:
        return txt


def html_string_to_pdf(html_str, url_path=None, logger=None):
    """
    Generate a pdf:str and associated metadata (title, toc) from html string content
    """
    if logger:
        time0 = time.time()
    meta = ''
    result = io.BytesIO()
    try:
        html = io.StringIO(html_str) 
        meta = pisa.pisaDocument(src=html,
                                 dest=result,
                                 path=url_path)
    except:
        print("Error: unable to create the PDF")
    if meta and logger:
        time1 = time.time()
        logger.info(f'Convert html to pdf took: {time1 - time0} secs')
    pdf = result.getvalue()
    return meta, pdf






'''
def extract_docx_original(self, logger):
    """Extract from docx filetype.
    
    'text': excerpts.paragraphs
    """
    try:
        excerpts = docx.Document(self.filepath)
    except:
        logger.info("TypeError: document not of type `.docx`")
        return None
    txt = [par.text for par in excerpts.paragraphs]

    record['title'] = excerpts.paragraphs[0].text     #<<< get_title(txt)
    record['length_lines'] = len( clean_text(txt).split('.') )
    record['text'] = txt
    return record
'''

def extract_docx(self, logger):
    """Extract from docx using `extract_pdf`
    
    'text': excerpts.paragraphs
    """
    try:
        excerpts = docx.Document(self.filepath)
    except:
        logger.info("TypeError: document not of type `.docx`")
        return None
    txt_lst = [par.text for par in excerpts.paragraphs]
    txt = ''.join(txt_lst)

    tmp_file = self.filepath.parent / (self.filepath.stem + '.pdf')
    html_string_to_pdf(html_str = txt, 
                       output = tmp_file
                      )
    self.filepath = tmp_file
    record = extract_pdf(self, logger)
    return record

'''
def extract_html_original(self, logger):
    """Extract for html filetype.

    Use the optimal module, first, this is typically 
    PyMuPDF(fitz) for speed of extraction.  Then 
    try pdfminer.six, which is more accurate,  if any 
    errors occur.
    
    'text': visible text
    """
    #ref: https://stackoverflow.com/questions/1936466/how-to-scrape-only-visible-webpage-text-with-beautifulsoup
    def tag_visible(element):
        if element.parent.name in ['style', 'script', 'head', 'title', 'meta', '[document]']:
            return False
        if isinstance(element, Comment):
            return False
        return True

    def text_from_html(soup):
        texts = soup.findAll(text=True)
        visible_texts = filter(tag_visible, texts)  
        return u" ".join(t.strip() for t in visible_texts)

    with open(self.filepath.__str__(), 'r') as f:
        html = f.readlines()
    clean_html = ('').join(html).replace('\n','')
    try:
        excerpts = bs4.BeautifulSoup(clean_html, 'html.parser')
    except:
        logger.info("TypeError: document not of type `.html`")
        return None
    txt = text_from_html(excerpts)
    record['title'] = excerpts.find('title').text
    record['text'] = txt
    return record
'''
    

def extract_html(self, logger):
    """Extract from html using `extract_pdf`
    TODO: move file discrimination outside of function
    """
    url_path = None
    if type(self.filepath) == UniformResourceLocator and self.filetype in ['.html','.pdf']:
        html = self.file_str
        url_path = str(self.url)
    elif type(self.filepath) == PosixPath:
        with open(self.filepath.__str__(), 'r') as f:
            html = f.read()
    else:
        raise TypeError
    meta, pdf = html_string_to_pdf(html_str=html, 
                                   url_path=url_path,
                                   logger=logger)
    #record = extract_pdf(pdf, logger)
    newPdf = pypdf.PdfReader(io.BytesIO(pdf))
    newPdf.pages.__len__()

    record = meta.meta
    record['toc'] = meta.toc
    return record


def extract_pdf(pdf, logger):
    """Extract from pdf filetype.
    
    'text': pages
    """
    def get_metadata(pdf):
        author = None
        '''
        if type(self.filepath) == UniformResourceLocator:
            with fitz.open(stream=self.file_str) as doc:
                tmp = doc.metadata
                tmp['page_count'] = len(doc)
        elif type(self.filepath) == PosixPath:
            with fitz.open(filename=self.filepath.__str__()) as doc:
                tmp = doc.metadata
                tmp['page_count'] = len(doc)
        '''
        if pdf:
            with fitz.open(stream=pdf) as doc:
                tmp = doc.metadata
                tmp['page_count'] = len(doc)
        else:
            raise TypeError
        result = {'author': tmp['author'], 
                  'subject': tmp['subject'], 
                  'keywords': tmp['keywords'],
                  'page_count': tmp['page_count'], 
                  'date': datetime.datetime.strptime(tmp['creationDate'].split('D:')[1][:8], "%Y%m%d").date()
                  }
        return result
    
    def get_title(self):
        title = None
        try:
            with timeout(seconds=5):
                title = pdftitle.get_title_from_file(self.filepath.__str__())
        except Exception:
            logger.info("`pdftitle` module threw error")
            pass

        if not title:
            with timeout(seconds=5):
                with open(self.filepath.__str__(), 'rb') as f:
                    pdfReader = pypdf.PdfReader(f)
                    first_str = len(excerpts[0])
                    if pdfReader.metadata['/Title']:
                        title = pdfReader.metadata['/Title']
                    elif first_str>0 and first_str<100:
                        title = first_str
                    else:
                        pass
        return title

    def get_toc(self):
        outline = None
        pdf = Pdf.open(self.filepath.__str__())         #ref[https://stackoverflow.com/questions/48157194/how-can-i-extract-the-toc-with-pypdf2]
        if len(pdf.open_outline().root) > 1:
            outline = []
            for item in pdf.open_outline().root:
                outline.append((item.title, item.page_location))
        if not outline and applyPyMuPDF:
            with fitz.open(self.filepath.__str__()) as doc:         #stream=io.BytesIO(result.getvalue())
                tmp = doc.get_toc()
            outline = tmp if tmp != [] else None
        if not outline:
            with open(self.filepath.__str__(), 'rb') as fp:
                parser = PDFParser(fp)                          #PDFParser(io.BytesIO(result.getvalue()))
                document = PDFDocument(parser)
                try:
                    outline = list(document.get_outlines())
                except:
                    pass
        return outline

    def get_raw_text(self, number_of_pages_to_extract_text):
        """Get raw text from pdf.

        Ensure only a limited number of pages are extracted.
        """
        MAX_TIME_SEC = 10
        raw_text = ''

        try:
            with fitz.open(self.filepath.__str__() ) as doc:
                pg_idx = 0
                for page in doc:
                    if pg_idx <= number_of_pages_to_extract_text:
                        raw_text += page.get_text()
                        pg_idx += 1
        except Exception:
            pass    

        if raw_text == '':
            try:
                with timeout(seconds=MAX_TIME_SEC):
                    raw_text = pdf_extract_text(pdf_file = self.filepath.__str__(),
                                                maxpages = number_of_pages_to_extract_text
                                                )
            except Exception:
                logger.info(f'{self.filepath.__str__()} took more than {MAX_TIME_SEC}sec to extract text')

            if not raw_text=='':
                raw_text = pdf_extract_text(pdf_file = self.filepath.__str__(),
                                            maxpages = 1
                                            )
        return raw_text        


    #process raw data
    time0 = time.time()
    metadata_results = get_metadata(pdf)
    page_list_length = metadata_results['page_count'] - 1
    if MAX_PAGE_EXTRACT: 
        number_of_pages_to_extract_text = page_list_length if page_list_length <= MAX_PAGE_EXTRACT else MAX_PAGE_EXTRACT
    else:
        number_of_pages_to_extract_text = page_list_length
    raw_text = get_raw_text(self, number_of_pages_to_extract_text)
    excerpts = clean_text(raw_text)

    #create record
    record['author'] = metadata_results['author']
    record['subject'] = metadata_results['subject']
    record['keywords'] = metadata_results['keywords']
    record['date'] = metadata_results['date']
    record['page_nos'] = metadata_results['page_count']
    record['title'] = get_title(self)
    record['toc'] = get_toc(self)
    record['body'] = excerpts

    time1 = time.time()
    logger.info(f'Convert html to pdf took: {time1 - time0} secs')

    return record


def extract_csv(self, logger):
    df = pd.read_csv(self.filepath.__str__(), 
                    nrows=1
                    )
    record['title'] = df.columns[0].replace('\n','')
    return record


def extract_xlsx(self, logger):
    sheets = pd.ExcelFile(self.filepath.__str__()).sheet_names
    df = pd.read_excel(self.filepath.__str__(), 
                        sheet_name=sheets[0], 
                        nrows=1
                        )

    record['title'] = df.columns[0].replace('\n','')
    return record